from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# --- Styles ---
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# --- Title ---
title = doc.add_heading('Backend Enhancement Tasks', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Subtitle
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('FutureIntern Platform — Spring 2026')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph()

# --- Tasks ---
tasks = [
    {
        "num": 1,
        "title": "Internship Lifecycle, Logging & Admin Analytics",
        "timeline": "Week 2 → Week 3",
        "desc": "Provide the backend endpoints the frontend team needs for the Admin Dashboard. Includes admin analytics endpoints returning stats like total users, applications per internship, and acceptance rates, structured logging for every API request (who, what, when, from which IP), an audit trail that records sensitive actions (login attempts, password changes, role changes, application status updates) in a history table, automatic filtering of expired internships so they disappear from the dashboard once their deadline passes, and an apply redirect that returns the real company application form URL (from the CSV data) so students apply through the official company portal. Ready before frontend starts Admin Dashboard in Week 2."
    },
    {
        "num": 2,
        "title": "CV Builder API",
        "timeline": "Week 3 → Week 4",
        "desc": "Provide the backend API the frontend team needs for the interactive CV Builder. Includes CRUD endpoints for CV sections (education, experience, skills, projects), PDF export so students can download their CV, file upload security to prevent malicious files, and data access protection so each user can only view and edit their own CV. Ready before frontend starts CV Builder in Week 3."
    },
    {
        "num": 3,
        "title": "Comprehensive Security Hardening",
        "timeline": "Week 4 → Week 6",
        "desc": "Protect the backend from common attacks and enforce access control while frontend focuses on Theme and Home Page (no backend dependency). Includes JWT token blacklisting & refresh rotation so stolen tokens can be revoked, role-based access control (RBAC) so students can't access admin endpoints, SQL injection / XSS / CSRF prevention, input validation with Marshmallow schemas to reject malformed data, rate limiting & account lockout to block brute-force attacks, Two-Factor Authentication (2FA) via email codes for an extra login layer, and browser back/forward button protection so users cannot navigate back to authenticated pages after logging out or forward to restricted pages without proper authorization."
    },
    {
        "num": 4,
        "title": "Authentication System — Password Reset, Email Verification & Google OAuth",
        "timeline": "Week 8 → Week 9",
        "desc": "Complete the authentication system with three features: (1) Secure forgot/reset password flow with time-limited tokens that expire after use, (2) Email verification on registration so only real email addresses can create accounts, with Flask-Mail SMTP integration and a fallback when email service is unavailable, and (3) Google OAuth integration allowing users to sign up and log in with their Google account, including account linking, secure token exchange, and auto-registration for new OAuth users."
    },
    {
        "num": 5,
        "title": "AI-Powered Chatbot Integration",
        "timeline": "Week 9 → Week 10",
        "desc": "Connect the existing chatbot to real AI APIs so it can give intelligent responses about internships. Includes integration with OpenAI / DeepSeek APIs, conversation history storage so the chatbot remembers context within a session, context-aware responses trained on the platform's internship data, secure API key management using environment variables, and rate limiting per user to control API costs."
    },
    {
        "num": 6,
        "title": "Mobile Application (React Native)",
        "timeline": "Week 5 → Week 10 (parallel)",
        "desc": "Build a cross-platform mobile application using React Native that connects to the existing Flask backend API. Starts early in Week 5 with project setup, navigation structure, and core screens, then expands throughout the semester. Includes all core features available on the web platform — browsing internships, applying, viewing dashboard, managing saved internships, and push notifications for application status updates. The mobile app reuses the same backend endpoints with token-based authentication, providing students a native mobile experience on both Android and iOS. Runs in parallel with other tasks — completely independent from the web frontend."
    },
]

for task in tasks:
    # Task heading
    heading = doc.add_heading(level=2)
    run = heading.add_run(f"Task {task['num']} — {task['title']}")

    # Timeline
    timeline_para = doc.add_paragraph()
    run_label = timeline_para.add_run("Timeline: ")
    run_label.bold = True
    run_label.font.size = Pt(11)
    run_value = timeline_para.add_run(task['timeline'])
    run_value.font.size = Pt(11)

    # Description
    desc_para = doc.add_paragraph(task['desc'])
    desc_para.paragraph_format.space_after = Pt(6)

    # Midterms notice after Task 3
    if task['num'] == 3:
        doc.add_paragraph()
        midterm = doc.add_paragraph()
        midterm.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = midterm.add_run('— Week 7: Midterms (No Work) —')
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(180, 0, 0)
        doc.add_paragraph()

# --- Timeline Summary Table ---
doc.add_paragraph()
doc.add_heading('Timeline Summary', level=1)

data = [
    ('Week 2–3', 'Task 1 — Admin Analytics, Logging & Internship Lifecycle'),
    ('Week 3–4', 'Task 2 — CV Builder API'),
    ('Week 4–6', 'Task 3 — Security Hardening'),
    ('Week 7', 'Midterms — No Work'),
    ('Week 8–9', 'Task 4 — Auth: Password Reset, Email & Google OAuth'),
    ('Week 5–10', 'Task 6 — Mobile App (parallel)'),
    ('Week 9–10', 'Task 5 — AI Chatbot'),
]

table = doc.add_table(rows=len(data) + 1, cols=2, style='Light Shading Accent 1')
table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Header
table.rows[0].cells[0].text = 'Week'
table.rows[0].cells[1].text = 'Activity'
for cell in table.rows[0].cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True

# Data rows
for i, (week, activity) in enumerate(data):
    table.rows[i + 1].cells[0].text = week
    table.rows[i + 1].cells[1].text = activity

# --- Milestones ---
doc.add_paragraph()
milestone = doc.add_paragraph()
run1 = milestone.add_run('First 3 tasks completed by: ')
run1.bold = True
milestone.add_run('End of Week 6')

milestone2 = doc.add_paragraph()
run2 = milestone2.add_run('Midterms break: ')
run2.bold = True
milestone2.add_run('Week 7')

milestone3 = doc.add_paragraph()
run3 = milestone3.add_run('All 6 tasks completed by: ')
run3.bold = True
milestone3.add_run('End of Week 10')

# --- Save ---
output_path = r"d:\SUT\Fall 2025\Project I\futureinternn\futureintern\Backend_Enhancement_Tasks.docx"
doc.save(output_path)
print(f"Document saved to: {output_path}")
